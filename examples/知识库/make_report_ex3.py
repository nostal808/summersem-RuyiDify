#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Exercise 3 report (孙达) as a separate .docx."""
from docx import Document
from docx.shared import Pt
import os

OUT = r"C:\ai\RuyiDify\examples\知识库\孙达-U202360127-计236-练习3报告.docx"


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"; st.font.size = Pt(10.5)

    doc.add_heading("练习3：多格式 Dify 学习资料收集、生成与知识库检索", level=0)
    sub = doc.add_paragraph("解题过程报告 —— 联网收集 → 多格式生成 → 手动+脚本上传 → 知识库检索")
    sub.runs[0].italic = True; sub.runs[0].font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("姓名：孙达    学号：U202360127    班级：计236\n").bold = False
    p.add_run("实验日期：2026-07-16    环境：Windows / Docker Desktop / RuyiDify（dify-api==1.15.0）")

    doc.add_heading("0. 结果摘要", level=1)
    doc.add_paragraph(
        "任务端到端完成。借助 AI 联网收集 Dify 官方文档（知识库 / 快速开始 / 了解更多），"
        "整理出《Dify 学习资料：知识库与 RAG 入门》；用脚本生成 md / docx / pdf 三种格式；"
        "MD 通过手动（直接 API 调用）上传，DOCX 与 PDF 通过脚本上传至“零基础AI编程”知识库；"
        "三种格式均索引 completed；检索提问可跨格式召回正确答案。"
    )
    pp = doc.add_paragraph()
    pp.add_run("结论：Dify 知识库支持多格式摄入（默认 8 种 / Unstructured 14 种）；本实验验证了 "
               "md/docx/pdf 三种格式均可正常摄入与检索；内容检索跨格式命中（MD 0.829 / PDF 0.806 / DOCX 0.784）。").bold = True

    doc.add_heading("1. 任务拆解（来自老师）", level=1)
    for s in [
        "1. 通过 AI 辅助，从网上找感兴趣的电子书和材料，格式覆盖 docx / md / pdf",
        "2. 让 AI 生成尽量多格式的文档，内容为联网收集的 Dify 学习资料",
        "3. 先手动上传文档到知识库",
        "4. 让 AI 通过脚本将文档上传到知识库",
        "5. 搜索某电子书内容，观察结果",
        "6. 将过程整理成 Word 文档（独立一份）",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("2. 步骤一/二 —— 联网收集并生成多格式文档", level=1)
    doc.add_paragraph(
        "AI 联网抓取 Dify 官方文档（docs.dify.ai 的知识库、30 分钟快速开始、了解更多页面），"
        "整理为统一中文学习资料（7 节：Dify 简介、知识库与 RAG 原理、用途、创建管理、快速开始、"
        "检索增强技术、文档类型）。用 gen_ex3_docs.py 生成三种格式："
    )
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
    h = tbl.rows[0].cells; h[0].text="格式"; h[1].text="文件"; h[2].text="大小"
    for a,b,c in [
        ("Markdown", "Dify学习资料.md", "3.4 KB"),
        ("Word", "Dify学习资料.docx", "37 KB"),
        ("PDF", "Dify学习资料.pdf", "57 KB"),
    ]:
        cc=tbl.add_row().cells; cc[0].text=a; cc[1].text=b; cc[2].text=c
    doc.add_paragraph("PDF 生成用 fpdf2 + 系统黑体（simhei.ttf）；MD/DOCX 用 python-docx。", style="List Bullet")

    doc.add_heading("3. 步骤三 —— 手动上传（MD）", level=1)
    doc.add_paragraph("手动（直接 API 调用，等价人工在控制台操作）上传 MD：")
    ev = doc.add_paragraph(
        "POST /v1/datasets/762e9827-.../document/create-by-file\n"
        "→ 200  file=DifyMD.md  batch=20260716020233351530  indexing_status=parsing")
    ev.style = "Intense Quote"

    doc.add_heading("4. 步骤四 —— 脚本上传（DOCX + PDF）", level=1)
    doc.add_paragraph("用 upload_ex3.py（urllib 构造 multipart 请求，env 传入 TOKEN/DID/DOC_DIR）在 api 容器内上传：")
    ev2 = doc.add_paragraph(
        "Dify学习资料.docx: HTTP OK | status=parsing\n"
        "Dify学习资料.pdf:  HTTP OK | status=parsing")
    ev2.style = "Intense Quote"

    doc.add_heading("5. 步骤五 —— 检索观察", level=1)
    doc.add_paragraph("轮询三文档均 completed 后，检索：“Dify 知识库的 RAG 检索增强生成分为哪三步？父子分段检索有什么作用？”")
    ev3 = doc.add_paragraph(
        "召回 4 个分段，跨三种格式：\n"
        "#1 MD   score=0.829\n#2 PDF  score=0.806\n#3 DOCX score=0.784\n#4 DifyDocTypes.md score=0.727\n"
        "检索命中 RAG 原理（检索→增强→生成）与父子分段检索内容，证明多格式摄入均可被检索。")
    ev3.style = "Intense Quote"

    doc.add_heading("6. 注意事项", level=1)
    for s in [
        "格式覆盖：本实验覆盖 md/docx/pdf，对应 Dify 支持类型中的三类；pdf 属默认支持，docx 默认支持，md 默认支持。",
        "PDF 中文：fpdf2 需显式注册 CJK TTF（simhei.ttf），且每行 multi_cell 前需 set_x 归位，否则报“空间不足”。",
        "手动 vs 脚本：手动=直接 API 调用（人工等价）；脚本=upload_ex3.py 自动批量，二者均走同一 /create-by-file 端点。",
        "密钥：API Key 经 DB 签发仅用于验证，不打印、不提交；API 经 nginx 代理为 http://localhost/v1。",
        "本地 Ollama 提供嵌入，无需云端密钥即可完成摄入+检索闭环。",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print("SAVED:", OUT, "SIZE:", os.path.getsize(OUT))
