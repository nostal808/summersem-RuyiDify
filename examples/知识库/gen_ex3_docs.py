#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Exercise 3: 生成多格式 Dify 学习资料文档（MD / DOCX / PDF）。
内容基于 Dify 官方文档（docs.dify.ai）联网收集整理，2026-07-16。
输出目录：examples/知识库/ex3_docs/
"""
import os

OUTDIR = r"C:\ai\RuyiDify\examples\知识库\ex3_docs"
os.makedirs(OUTDIR, exist_ok=True)

# ---- 共享正文（真实整理自官方文档）----
TITLE = "Dify 学习资料：知识库与 RAG 入门"
META = "整理来源：Dify 官方文档 docs.dify.ai（知识库 / 快速开始 / 了解更多）｜整理日期：2026-07-16"

SECTIONS = [
    ("1. Dify 是什么",
     "Dify 是一个开源的 AI 应用开发平台，用于构建智能体（Agent）、智能体工作流（agentic "
     "workflow）和聊天机器人，并基于你自己的数据进行增强，最终以 Web 应用或 API 形式发布。"
     "名称 Dify 源自 Do It For You。"),
    ("2. 知识库与 RAG 原理",
     "知识库（Knowledge）是你可以接入 AI 应用的自有数据集合。通过检索增强生成（RAG，"
     "Retrieval-Augmented Generation），让大模型以你的专属知识作为额外事实来源，使回答更准确、"
     "相关、更少幻觉。流程分三步：\n"
     "  (1) 检索（Retrieval）：用户提问时，系统先从知识库召回最相关内容；\n"
     "  (2) 增强（Augmented）：检索结果拼接用户原问题，作为增强上下文送入 LLM；\n"
     "  (3) 生成（Generation）：LLM 基于该上下文生成更精确的回答。\n"
     "知识库以 Dataset 形式存储与管理，可创建多个知识库，分别适配不同领域与数据源，按需接入应用。"),
    ("3. 知识库常见用途",
     "- 客服聊天机器人：基于产品文档、FAQ、排障指南给出准确答复；\n"
     "- 企业内部知识门户：员工快速检索公司制度与流程；\n"
     "- 内容生成工具：基于背景素材生成报告、文章、邮件；\n"
     "- 研究分析应用：从论文、市场报告、法律文档等仓库检索并摘要。"),
    ("4. 创建与管理知识库",
     "- 创建即用知识库：导入数据、定义处理规则，Dify 自动完成后续步骤，适合入门；\n"
     "- 自定义知识库（知识流水线）：用自定义步骤与多种集成编排更复杂的数据处理工作流；\n"
     "- 连接外部知识库：通过 API 直接同步外部知识库，免迁移复用既有数据。\n"
     "管理侧：维护文档与分段、测试检索效果、用元数据增强检索、随时调整索引方式/"
     "嵌入模型/检索策略。"),
    ("5. 快速开始（以工作流为例）",
     "官方 30 分钟快速开始用 Dify Cloud 构建一个“多平台内容生成器”工作流：从空白创建工作流，"
     "用 User Input 节点收集草稿文本、参考文件、语气、目标平台、语言等输入；通过大模型节点按指定"
     "语气生成内容；用迭代/模板节点输出多平台文案。核心思想：用清晰的节点与变量编排复杂 AI 应用，"
     "比裸调 API 更快更省力。自托管实例步骤一致。"),
    ("6. 检索增强技术（官方博客要点）",
     "- 自定义元数据过滤（v1.1.0）：为文档加元数据，支持基于过滤的检索，提升精度；\n"
     "- 父子分段检索（v0.15.0）：父分段保上下文、子分段保精度，召回更优；\n"
     "- 混合检索 + 重排序（Hybrid Search + Rerank）：结合向量与关键词检索并排序，提升 RAG 准确率；\n"
     "- 文本嵌入基础：嵌入模型把文本映射为向量，是检索与相似度计算的基础。"),
    ("7. 文档类型与接入（与本实验关联）",
     "Dify 知识库支持多格式文档摄入：默认模式 8 种（txt/md/pdf/html/docx/csv/xlsx/epub），"
     "Unstructured 模式 14 种（新增 doc/msg/eml/ppt/pptx/xml），另有 Notion 与 Website 两类数据源。"
     "本实验生成 md/docx/pdf 三种格式，验证多格式摄入与检索。"),
]

def md_text() -> str:
    lines = [f"# {TITLE}", "", f"> {META}", ""]
    for h, b in SECTIONS:
        lines.append(f"## {h}")
        lines.append("")
        lines.append(b)
        lines.append("")
    return "\n".join(lines)

def write_md():
    p = os.path.join(OUTDIR, "Dify学习资料.md")
    with open(p, "w", encoding="utf-8") as f:
        f.write(md_text())
    return p

def write_docx():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Microsoft YaHei"; st.font.size = Pt(11)
    doc.add_heading(TITLE, level=0)
    m = doc.add_paragraph(); m.add_run(META).italic = True
    for h, b in SECTIONS:
        doc.add_heading(h, level=1)
        for para in b.split("\n"):
            doc.add_paragraph(para)
    p = os.path.join(OUTDIR, "Dify学习资料.docx")
    doc.save(p)
    return p

def _san(s: str) -> str:
    # fpdf2 行断点处理全角标点会报错，PDF 内统一替换为半角
    rep = {"｜": "|", "（": "(", "）": ")", "：": ": ", "，": ", ", "。": ". ", "、": ", ",
           "“": '"', "”": '"', "《": "<", "》": ">", "／": "/", "　": " "}
    for k, v in rep.items():
        s = s.replace(k, v)
    return s


def write_pdf():
    from fpdf import FPDF
    pdf = FPDF(format="A4")
    pdf.add_page()
    font = "Helvetica"
    try:
        pdf.add_font("cn", "", r"C:\Windows\Fonts\simhei.ttf", uni=True)
        font = "cn"
    except Exception:
        font = "Helvetica"
    pdf.set_font(font, "", 16)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 9, _san(TITLE))
    pdf.set_font(font, "", 9)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 5, _san(META))
    pdf.ln(2)
    for h, b in SECTIONS:
        pdf.set_font(font, "", 12)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 7, _san(h))
        pdf.set_font(font, "", 10)
        for para in b.split("\n"):
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, _san(para))
        pdf.ln(1)
    p = os.path.join(OUTDIR, "Dify学习资料.pdf")
    pdf.output(p)
    return p

if __name__ == "__main__":
    for fn in (write_md, write_docx, write_pdf):
        print("WROTE:", fn())
